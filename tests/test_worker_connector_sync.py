"""Tests for IM-3 + Task 4 — ``run_connector_sync_pipeline`` enumerates
the canonical ``topology.connectors`` block and wires the production
``ConnectorPipeline`` against each cc_pair.

Task 4 of the connector canonical-collapse refactor (Phase 1) moves
ingest enumeration off the legacy top-level ``connectors:`` list onto
``topology.connectors`` read through the overlay-aware merged mapping
(``ConnectorSyncDeps.config_mapping_fn``). The overloaded legacy
``entry["name"]`` splits into three canonical values: ``kind`` (plugin
resolution), the cc_pair ``name`` (routing / chunk-writer key), and
``config`` (the connector_specific_config mapping).

Covers:
  - disabled short-circuit returns zero counters without touching the
    config / DB / bronze paths;
  - end-to-end run against a real Obsidian vault + passthrough extractor
    enumerated from ``topology.connectors`` + a cc_pair, indexes both
    markdown files and lands them in the cc_pair-named collection;
  - per-connector failure is logged and the loop continues — sibling
    connectors still report their own counters;
  - a connector with zero cc_pairs is skipped (no collection target);
  - the ``connector_enabled`` predicate is consulted per entry — a
    registered kind gated OFF is skipped, a flagless sibling still runs.

Sabotage-proof (executed by the agent, recorded for the reader): in
``run_connector_sync_pipeline`` comment out the ``pipeline.run_batch(...)``
line inside ``_run_one_connector_batch``; re-run
``test_runs_configured_obsidian_pipeline`` — the ``synced == 2``
assertion fails (the counters stay at zero because no item was
processed). Restore the call; the test passes again.
"""

from __future__ import annotations

import logging
import sqlite3
from pathlib import Path
from typing import Any

import pytest

from kairix.core.connectors.cc_pair import create_cc_pair
from kairix.core.db.schema import create_schema
from kairix.core.protocols import BronzeRef, DocMetadata, ExtractedDocument, Page
from kairix.worker import (
    ConnectorSyncDeps,
    ConnectorSyncResult,
    ConnectorSyncRuntime,
    build_worker_silver_processor,
    run_connector_sync_pipeline,
)
from tests.fakes import FakeSourceConnector

# NOTE: no module-level ``pytestmark`` — a module-level ``pytest.mark.unit``
# STACKS with the per-function ``@pytest.mark.integration`` markers below,
# so the real-SQLite / real-connector integration tests would wrongly also
# carry the unit marker and run in the unit gate. Each test instead declares
# its own marker individually: pure-logic tests are ``@pytest.mark.unit``;
# tests that drive a real SQLite DB + real connector pipeline are
# ``@pytest.mark.integration`` ONLY.


def _seed_cc_pair_row(db_path: Path, *, cc_pair_name: str) -> None:
    """Pre-register a ``topology_cc_pairs`` row (connector + cc_pair) so the
    worker's name→id lookup resolves and the lifetime counter can be bumped.

    Mirrors the real flow: the operator applies their topology (which lands a
    ``topology_cc_pairs`` row) BEFORE the worker syncs. ``run_connector_sync_pipeline``
    itself does not persist config cc_pairs into the table — that's the
    topology applier's job — so the test seeds the row directly.
    """
    db = sqlite3.connect(str(db_path))
    try:
        create_schema(db)
        now = "2026-06-22T00:00:00Z"
        cur = db.execute(
            "INSERT INTO topology_connectors "
            "(kind, name, connector_specific_config, default_sensitivity, created_at, updated_at) "
            "VALUES ('obsidian', 'seed-connector', '{}', 'internal', ?, ?)",
            (now, now),
        )
        connector_id = cur.lastrowid
        assert connector_id is not None
        create_cc_pair(db, connector_id=int(connector_id), credential_id=None, name=cc_pair_name)
        db.commit()
    finally:
        db.close()


def _total_docs_indexed(db_path: Path, *, cc_pair_name: str) -> int:
    db = sqlite3.connect(str(db_path))
    try:
        row = db.execute(
            "SELECT total_docs_indexed FROM topology_cc_pairs WHERE name = ?",
            (cc_pair_name,),
        ).fetchone()
    finally:
        db.close()
    assert row is not None, f"no topology_cc_pairs row for {cc_pair_name!r}"
    return int(row[0])


def _reset_connector_cursor(db_path: Path) -> None:
    """Clear persisted connector cursors so the next sync is a cold reconcile.

    Obsidian's cursor is an ISO mtime high-water-mark; a second short-lived
    run within the same wall-clock window would otherwise surface zero events
    (no live watchdog observer ran between the two batches). Clearing the
    cursor forces the reconciler to re-emit the vault as a cold start, which
    deterministically re-processes the notes so the increment behaviour can
    be asserted without depending on filesystem-event timing.
    """
    db = sqlite3.connect(str(db_path))
    try:
        db.execute("DELETE FROM connector_cursors")
        db.commit()
    finally:
        db.close()


def _no_db_factory() -> sqlite3.Connection:
    """Sentinel db_factory that asserts when the short-circuit path is bypassed."""
    raise AssertionError("db_factory must not be invoked on the short-circuit path")


class _CloseTrackingConnector(FakeSourceConnector):
    """No-change connector that records worker-owned lifecycle closure."""

    def __init__(self) -> None:
        super().__init__()
        self.close_calls = 0

    def close(self) -> None:
        self.close_calls += 1


def _exercise_long_lived_obsidian_runtime(
    vault: Path,
    db_path: Path,
    bronze_root: Path,
    result_sink: Any,
) -> None:
    """Run the real watcher lifecycle in an isolated worker process.

    macOS's FSEvents and lxml native extensions are unsafe when loaded into
    the same pytest process.  Production also runs the watcher in a dedicated
    worker process, so the integration test mirrors that boundary and reports
    only the observable per-tick sync counts to its parent.
    """
    from watchdog.observers.polling import PollingObserver

    from kairix.connectors.obsidian import ObsidianConnector
    from kairix.connectors.obsidian.watcher import WatchdogSource

    def _resolve_obsidian(_kind: str) -> Any:
        def _build(config: dict[str, Any]) -> ObsidianConnector:
            def _polling_watcher(root: Path) -> WatchdogSource:
                return WatchdogSource(root, observer_factory=PollingObserver)

            return ObsidianConnector(
                vault_root=Path(config["vault_root"]),
                watcher_factory=_polling_watcher,
            )

        return _build

    runtime = ConnectorSyncRuntime(
        deps=ConnectorSyncDeps(
            disabled_fn=lambda: False,
            config_mapping_fn=lambda: _obsidian_topology(vault),
            db_factory=lambda: sqlite3.connect(str(db_path)),
            bronze_root_resolver=lambda: bronze_root,
        ),
        connector_factory_resolver=_resolve_obsidian,
    )
    try:
        result_sink.put([runtime().synced for _ in range(10)])
    finally:
        runtime.close()


def _obsidian_topology(vault: Path, *, cc_pair_name: str = "obsidian-personal") -> dict[str, Any]:
    """Build a minimal merged mapping with one obsidian connector + cc_pair.

    Mirrors the canonical ``topology.connectors`` / ``topology.cc_pairs``
    shape the setup wizard writes — the connector carries ``kind`` +
    ``connector_specific_config``, and a single cc_pair binds it to a
    routing name (the chunk-writer collection key).
    """
    return {
        "topology": {
            "connectors": [
                {
                    "id": "obsidian-conn",
                    "kind": "obsidian",
                    "name": "Personal Obsidian Vault",
                    "extractor": "passthrough",
                    "connector_specific_config": {"vault_root": str(vault)},
                }
            ],
            "cc_pairs": [
                {
                    "id": "obsidian-pair",
                    "connector": "obsidian-conn",
                    "credential": None,
                    "name": cc_pair_name,
                }
            ],
        }
    }


@pytest.mark.integration
def test_connector_runtime_reuses_one_instance_for_more_than_inotify_limit_and_closes_it(tmp_path: Path) -> None:
    """A worker lifetime owns one connector across ticks and closes it once.

    The 129 ticks cross Linux's common ``max_user_instances=128`` boundary.
    Reconstructing one watcher-backed connector per tick leaks the first 128
    observers and fails production on tick 129.  The runtime must resolve the
    connector once, reuse it, and close it deterministically at shutdown.
    """
    connector = _CloseTrackingConnector()
    factory_calls = 0

    def _resolve(_kind: str) -> Any:
        def _factory(_config: dict[str, Any]) -> _CloseTrackingConnector:
            nonlocal factory_calls
            factory_calls += 1
            return connector

        return _factory

    vault = tmp_path / "vault"
    vault.mkdir()
    runtime = ConnectorSyncRuntime(
        deps=ConnectorSyncDeps(
            disabled_fn=lambda: False,
            config_mapping_fn=lambda: _obsidian_topology(vault),
            db_factory=lambda: sqlite3.connect(str(tmp_path / "index.sqlite")),
            bronze_root_resolver=lambda: tmp_path / "bronze",
        ),
        connector_factory_resolver=_resolve,
    )

    for _ in range(129):
        runtime()

    assert factory_calls == 1
    assert connector.close_calls == 0

    runtime.close()
    runtime.close()

    assert connector.close_calls == 1


@pytest.mark.integration
def test_long_lived_runtime_reaches_obsidian_periodic_reconciliation(tmp_path: Path) -> None:
    """The real Obsidian connector retains its ten-tick reconciliation cadence.

    A newly constructed connector cold-scans on every call.  One worker-owned
    connector instead cold-scans once, remains quiet for ticks two through
    nine, and runs its configured periodic reconciliation on tick ten.  This
    exercises the real filesystem watcher, reconciler, pipeline, and SQLite
    cursor rather than a lifecycle stand-in.
    """
    vault = tmp_path / "vault"
    vault.mkdir()
    (vault / "alpha.md").write_text("# Alpha\n\nLong-lived runtime reconciliation.\n", encoding="utf-8")
    db_path = tmp_path / "index.sqlite"
    from multiprocessing import get_context

    context = get_context("spawn")
    result_sink = context.Queue()
    process = context.Process(
        target=_exercise_long_lived_obsidian_runtime,
        args=(vault, db_path, tmp_path / "bronze", result_sink),
    )
    process.start()
    process.join(timeout=30)

    assert process.exitcode == 0
    synced = result_sink.get(timeout=1)
    result_sink.close()
    process.close()
    assert synced[0] == 1
    assert synced[1:9] == [0] * 8
    assert synced[9] == 1


@pytest.mark.integration
def test_sync_pipeline_threads_extractor_chain_configs_to_chain_members(
    tmp_path: Path,
    caplog: pytest.LogCaptureFixture,
) -> None:
    """The public sync pipeline must carry per-chain extractor config.

    Gotenberg claims DOCX only when ``extractor_chain_configs.gotenberg``
    reaches the registry with ``include_docx: true``. This test uses a real
    Obsidian connector over a DOCX file, points Gotenberg at an invalid URL so
    it raises immediately, then expects the chain to fall back to the real
    DOCX extractor and index the file.

    If the worker drops ``extractor_chain_configs`` while building the entry,
    Gotenberg never claims DOCX, no warning is logged, and this regression pin
    fails without importing private worker helpers (F5).
    """
    from docx import Document

    vault = tmp_path / "vault"
    vault.mkdir()
    doc_path = vault / "strategy.docx"
    document = Document()
    document.add_heading("Agent Exchange Strategy", level=1)
    document.add_paragraph(
        "This document is long enough for the heading-aware fallback extractor to pass quality checks. " * 4
    )
    document.save(doc_path)

    mapping = {
        "topology": {
            "connectors": [
                {
                    "id": "obsidian-conn",
                    "kind": "obsidian",
                    "name": "Personal Obsidian Vault",
                    "extractor_chain": ["gotenberg", "docx"],
                    "extractor_chain_configs": {
                        "gotenberg": {
                            "config": {
                                "gotenberg_url": "not-a-url",
                                "timeout_s": 0.01,
                                "include_docx": True,
                            }
                        }
                    },
                    "connector_specific_config": {
                        "vault_root": str(vault),
                        "collections": [{"name": "docx", "path": ".", "glob": "**/*.docx"}],
                    },
                }
            ],
            "cc_pairs": [
                {
                    "id": "obsidian-pair",
                    "connector": "obsidian-conn",
                    "credential": None,
                    "name": "obsidian-personal",
                }
            ],
        }
    }
    db_path = tmp_path / "index.sqlite"
    deps = ConnectorSyncDeps(
        disabled_fn=lambda: False,
        config_mapping_fn=lambda: mapping,
        db_factory=lambda: sqlite3.connect(str(db_path)),
        bronze_root_resolver=lambda: tmp_path / "bronze",
    )

    with caplog.at_level(logging.WARNING, logger="kairix.core.connectors.escalation"):
        result = run_connector_sync_pipeline(deps)

    assert result.synced == 1
    assert result.failed == 0
    assert any("gotenberg.extract raised RuntimeError" in record.getMessage() for record in caplog.records), (
        "gotenberg must claim DOCX via extractor_chain_configs before DOCX fallback indexes it"
    )


@pytest.mark.unit
def test_disabled_short_circuits(tmp_path: Path) -> None:
    """When ``deps.disabled_fn`` returns True, ``run_connector_sync_pipeline``
    returns a zero-counter :class:`ConnectorSyncResult` and never touches
    the DB / config path.

    Sabotage proof: change the early-return body to
    ``return ConnectorSyncResult(synced=1, ...)`` and the
    ``result.synced == 0`` assertion fails. Restored, the test passes.
    """
    deps = ConnectorSyncDeps(
        disabled_fn=lambda: True,
        config_mapping_fn=dict,
        db_factory=_no_db_factory,
        bronze_root_resolver=lambda: tmp_path / "bronze",
    )

    result = run_connector_sync_pipeline(deps)

    assert result == ConnectorSyncResult(synced=0, failed=0, dead_letter_added=0)


@pytest.mark.integration
def test_runs_configured_obsidian_pipeline(tmp_path: Path) -> None:
    """A real vault with two markdown files + a ``topology.connectors``
    block (one obsidian connector + one cc_pair) drives the full
    ``ConnectorPipeline`` and indexes both items.

    Proves the canonical-topology ingest redirect (Task 4): ingest
    enumerates ``topology.connectors`` from the merged mapping, NOT the
    legacy top-level ``connectors:`` list. This is the wizard-onboarding
    fix — the wizard writes ``topology.connectors`` and ingest now reads
    it.

    Uses real :class:`StreamingBronzeStore`, :class:`DefaultSilverProcessor`,
    :class:`CursorStore`, :class:`DeadLetterStore`, the in-process
    SQLite chunk-writer / entity-graph sink, and the real Obsidian
    connector + passthrough extractor resolved through the entry-point
    registry. No fakes at this seam — F47-clean.

    Sabotage proof (the brief's mandated one): comment out
    ``pipeline.run_batch(connector, extractor)`` inside
    ``_run_one_connector_batch``; the assertion ``result.synced == 2``
    fails because nothing flows through the pipeline. Restored, both
    notes are indexed and the test passes.
    """
    vault = tmp_path / "vault"
    vault.mkdir()
    (vault / "alpha.md").write_text("# Alpha\n\nFirst note body content.\n", encoding="utf-8")
    (vault / "beta.md").write_text("# Beta\n\nSecond note body content.\n", encoding="utf-8")

    db_path = tmp_path / "index.sqlite"

    deps = ConnectorSyncDeps(
        disabled_fn=lambda: False,
        config_mapping_fn=lambda: _obsidian_topology(vault),
        db_factory=lambda: sqlite3.connect(str(db_path)),
        bronze_root_resolver=lambda: tmp_path / "bronze",
    )

    result = run_connector_sync_pipeline(deps)

    assert result.synced == 2, (
        f"expected both notes synced through the pipeline; got {result}. "
        "fix: confirm ConnectorPipeline.run_batch processed each ChangeEvent and "
        "the writer.upsert call landed the chunks."
    )
    assert result.failed == 0
    assert result.dead_letter_added == 0

    db = sqlite3.connect(str(db_path))
    try:
        # #336 — wired-writer regression pin. Without
        # documents_media_writer=SqliteDocumentsMediaWriter(db) silver skips
        # the per-document row. Both notes should land one row each.
        (media_count,) = db.execute("SELECT COUNT(*) FROM documents_media").fetchone()
        # The chunk-writer routing key is the cc_pair name — chunks land in
        # the cc_pair-named collection, NOT the connector kind or the
        # connector name. Pins the D2 kind-vs-cc_pair-name split.
        collections = {row[0] for row in db.execute("SELECT DISTINCT collection FROM documents").fetchall()}
    finally:
        db.close()
    assert media_count == 2, (
        f"expected documents_media populated for both notes (#336 regression pin); got {media_count}. "
        "fix: confirm DefaultSilverProcessor is constructed with documents_media_writer=SqliteDocumentsMediaWriter(db) "
        "in kairix.worker._run_one_connector_batch and _build_reextract_components."
    )
    assert collections == {"obsidian-personal"}, (
        "chunks must land in the cc_pair-named collection (routing keys on cc_pair name, "
        f"not connector kind 'obsidian'); got {collections}."
    )


@pytest.mark.integration
def test_worker_silver_writes_document_pages_for_paged_extracts(tmp_path: Path) -> None:
    """The worker's Silver construction wires the page writer used by production sync.

    The factory-built connector pipeline already wires ``SqliteDocumentPagesWriter``;
    production worker sync uses ``build_worker_silver_processor``. A paged
    extraction must therefore write both ``documents_media`` and ``document_pages``
    rows through the worker helper, otherwise PDF/PPTX/DOCX page citations never
    become searchable in the live connector path.
    """
    db_path = tmp_path / "index.sqlite"
    db = sqlite3.connect(str(db_path))
    try:
        create_schema(db)
        silver = build_worker_silver_processor(db, read_flag=lambda _name: False)
        raw = BronzeRef(
            source_name="sharepoint",
            item_id="deck.pptx",
            raw_path=None,
            mime="application/vnd.openxmlformats-officedocument.presentationml.presentation",
            fetched_at="2026-07-20T00:00:00Z",
            content_hash="sha256-test-deck",
        )
        extracted = ExtractedDocument(
            markdown="# Slide 1\n\nBody\n\n# Slide 2\n\nMore body\n",
            pages=(
                Page(page_number=1, text="Slide 1 body", has_images=False),
                Page(page_number=2, text="Slide 2 body", has_images=True),
            ),
            images=(),
            metadata=DocMetadata(
                title="Deck",
                author=None,
                created_date=None,
                language="en",
                page_count=2,
            ),
            confidence=0.95,
        )

        silver.process(
            raw,
            extracted,
            source_uri="https://sharepoint.example/deck.pptx",
            source_modified_at="2026-07-20T00:00:00Z",
            sensitivity="internal",
            extractor_name="paged-test",
            extractor_version="v1",
        )

        media_count = db.execute("SELECT COUNT(*) FROM documents_media").fetchone()[0]
        page_rows = db.execute(
            "SELECT page_number, extracted_text, has_images FROM document_pages ORDER BY page_number"
        ).fetchall()
    finally:
        db.close()

    assert media_count == 1
    assert page_rows == [(1, "Slide 1 body", 0), (2, "Slide 2 body", 1)]


@pytest.mark.integration
def test_connector_sync_increments_cc_pair_total_docs_indexed(tmp_path: Path) -> None:
    """A successful batch bumps ``topology_cc_pairs.total_docs_indexed`` by the
    number of items processed — and a SECOND batch INCREMENTS (not overwrites).

    Wires the operator-facing counter that previously stayed at 0 forever, so
    ``kairix cc-pair list`` reports docs=0. The cc_pair row is pre-seeded (the
    operator applies topology before the worker syncs); each tick indexes new
    notes and the lifetime counter accrues.

    Proves increment (N then N+M) rather than recompute: the second batch's
    delta is ADDED to the first, not used to replace it.

    Sabotage proof: drop the ``total_docs_indexed = total_docs_indexed + ?``
    UPDATE in ``run_connector_sync_pipeline``'s loop → the counter stays 0 and
    the first assertion (``== 2``) fails. Restored, the counter accrues.
    """
    cc_pair_name = "obsidian-personal"
    vault = tmp_path / "vault"
    vault.mkdir()
    (vault / "alpha.md").write_text("# Alpha\n\nFirst note body content.\n", encoding="utf-8")
    (vault / "beta.md").write_text("# Beta\n\nSecond note body content.\n", encoding="utf-8")

    db_path = tmp_path / "index.sqlite"
    _seed_cc_pair_row(db_path, cc_pair_name=cc_pair_name)

    deps = ConnectorSyncDeps(
        disabled_fn=lambda: False,
        config_mapping_fn=lambda: _obsidian_topology(vault, cc_pair_name=cc_pair_name),
        db_factory=lambda: sqlite3.connect(str(db_path)),
        bronze_root_resolver=lambda: tmp_path / "bronze",
    )

    # First batch: indexes alpha + beta → counter = 2.
    first = run_connector_sync_pipeline(deps)
    assert first.synced == 2, f"expected both notes indexed in the first batch; got {first}"
    assert _total_docs_indexed(db_path, cc_pair_name=cc_pair_name) == 2

    # Second batch: reset the cursor so the reconciler re-emits the vault as a
    # cold start → the same two notes flow through again (processed == 2). The
    # lifetime counter must accrue to 4, proving INCREMENT rather than
    # recompute/overwrite (an overwrite would leave it at 2).
    _reset_connector_cursor(db_path)
    second = run_connector_sync_pipeline(deps)
    assert second.synced == 2, f"expected both notes re-processed on the second batch; got {second}"
    assert _total_docs_indexed(db_path, cc_pair_name=cc_pair_name) == 4, (
        "the counter must INCREMENT (2 + 2 = 4), not be overwritten by the second batch's delta"
    )


@pytest.mark.unit
def test_failing_connector_logged_and_loop_continues(
    tmp_path: Path,
    caplog: pytest.LogCaptureFixture,
) -> None:
    """With two configured connectors, the first one raises during
    resolution; the second one resolves and runs. The aggregated result
    must reflect only the second connector's counters and the worker
    must log the failure rather than crash.

    Uses a fictional ``does-not-exist`` kind for the first connector
    (forces ``resolve_connector`` to raise ``KeyError``) plus a valid
    obsidian connector against an empty vault for the second (zero items
    but the entry resolves cleanly). Each connector has a cc_pair so both
    are enumerated.

    Sabotage proof: remove the ``except Exception`` block inside the
    ``for entry in entries`` loop; the ``KeyError`` from the unknown
    connector propagates out of ``run_connector_sync_pipeline`` and the
    test fails with an unhandled error. Restored, the test passes
    because the per-entry try/except absorbs the failure.
    """
    empty_vault = tmp_path / "vault"
    empty_vault.mkdir()

    mapping: dict[str, Any] = {
        "topology": {
            "connectors": [
                {"id": "missing-conn", "kind": "does-not-exist", "name": "Missing"},
                {
                    "id": "obsidian-conn",
                    "kind": "obsidian",
                    "name": "Empty Vault",
                    "extractor": "passthrough",
                    "connector_specific_config": {"vault_root": str(empty_vault)},
                },
            ],
            "cc_pairs": [
                {"id": "missing-pair", "connector": "missing-conn", "credential": None, "name": "missing-cc"},
                {"id": "obsidian-pair", "connector": "obsidian-conn", "credential": None, "name": "obsidian-cc"},
            ],
        }
    }

    db_path = tmp_path / "index.sqlite"

    deps = ConnectorSyncDeps(
        disabled_fn=lambda: False,
        config_mapping_fn=lambda: mapping,
        db_factory=lambda: sqlite3.connect(str(db_path)),
        bronze_root_resolver=lambda: tmp_path / "bronze",
    )

    with caplog.at_level(logging.WARNING, logger="kairix.worker"):
        result = run_connector_sync_pipeline(deps)

    # First connector failed during resolve; second is a valid but empty
    # vault → zero items but no propagated exception. The aggregated
    # result reflects only the second connector's (zero) counters.
    assert result.synced == 0
    assert result.failed == 0
    assert result.dead_letter_added == 0

    failure_logs: list[str] = [
        rec.getMessage() for rec in caplog.records if "missing-cc" in rec.getMessage() and "failed" in rec.getMessage()
    ]
    assert failure_logs, (
        f"expected a warning naming the failing cc_pair; got {[r.getMessage() for r in caplog.records]}"
    )


@pytest.mark.unit
def test_no_topology_returns_zero_counter_no_op(tmp_path: Path) -> None:
    """No topology connectors → zero-counter result, no DB construction,
    no raise.

    Sabotage proof: change the early-return after the no-entries check
    to ``return ConnectorSyncResult(synced=9, ...)``; ``result.synced ==
    0`` fails. Restored, the test passes.
    """
    deps = ConnectorSyncDeps(
        disabled_fn=lambda: False,
        config_mapping_fn=dict,
        db_factory=_no_db_factory,
        bronze_root_resolver=lambda: tmp_path / "bronze",
    )

    result = run_connector_sync_pipeline(deps)

    assert result == ConnectorSyncResult(synced=0, failed=0, dead_letter_added=0)


@pytest.mark.unit
def test_zero_cc_pair_connector_is_skipped(tmp_path: Path) -> None:
    """A topology connector referenced by zero cc_pairs is not ingestable
    (no collection target) → skipped, no DB construction, zero counters.

    A connector with no cc_pair yields no entry, so the no-entries guard
    short-circuits before the db_factory is touched.

    Sabotage proof: yield a synthetic entry for a cc_pair-less connector
    (drop the ``connector.id in cc_pairs_by_connector_id`` guard) → the
    _no_db_factory sentinel fires when the loop builds the DB, failing
    this test.
    """
    vault = tmp_path / "vault"
    vault.mkdir()
    mapping: dict[str, Any] = {
        "topology": {
            "connectors": [
                {
                    "id": "lonely-conn",
                    "kind": "obsidian",
                    "name": "No cc_pair",
                    "connector_specific_config": {"vault_root": str(vault)},
                }
            ],
            # No cc_pairs referencing lonely-conn.
            "cc_pairs": [],
        }
    }

    deps = ConnectorSyncDeps(
        disabled_fn=lambda: False,
        config_mapping_fn=lambda: mapping,
        db_factory=_no_db_factory,
        bronze_root_resolver=lambda: tmp_path / "bronze",
    )

    result = run_connector_sync_pipeline(deps)

    assert result == ConnectorSyncResult(synced=0, failed=0, dead_letter_added=0)


@pytest.mark.integration
def test_connector_enabled_gates_per_entry_in_loop(
    tmp_path: Path,
    caplog: pytest.LogCaptureFixture,
) -> None:
    """The per-entry loop consults ``connector_enabled(entry["kind"], ...)``
    — a registered kind whose flag reads OFF is skipped, while a flagless
    sibling still runs.

    Two connectors: ``sharepoint`` (registered ``connector_sharepoint``
    flag, pinned OFF) and ``obsidian`` (flagless, always-on) over a real
    two-note vault. With the flag OFF, the sharepoint plugin never
    resolves (it would otherwise need a Graph credential); obsidian still
    indexes both notes.

    Sabotage proof: drop the ``connector_enabled`` skip in the loop — the
    sharepoint entry resolves + runs, but the ``gated off`` INFO log no
    longer fires, so the assertion below fails. Restored, the predicate
    skips sharepoint before resolution and the log fires.
    """
    vault = tmp_path / "vault"
    vault.mkdir()
    (vault / "alpha.md").write_text("# Alpha\n\nFirst note.\n", encoding="utf-8")
    (vault / "beta.md").write_text("# Beta\n\nSecond note.\n", encoding="utf-8")

    mapping: dict[str, Any] = {
        "topology": {
            "connectors": [
                {
                    "id": "sharepoint-conn",
                    "kind": "sharepoint",
                    "name": "Corp SharePoint",
                    "connector_specific_config": {},
                },
                {
                    "id": "obsidian-conn",
                    "kind": "obsidian",
                    "name": "Personal Vault",
                    "extractor": "passthrough",
                    "connector_specific_config": {"vault_root": str(vault)},
                },
            ],
            "cc_pairs": [
                {"id": "sp-pair", "connector": "sharepoint-conn", "credential": None, "name": "sharepoint-cc"},
                {"id": "ob-pair", "connector": "obsidian-conn", "credential": None, "name": "obsidian-cc"},
            ],
        }
    }

    db_path = tmp_path / "index.sqlite"

    def _flag_reader(_name: str) -> bool:
        # connector_sharepoint OFF (and any other registered kind defaults OFF).
        return False

    deps = ConnectorSyncDeps(
        disabled_fn=lambda: False,
        config_mapping_fn=lambda: mapping,
        flag_reader=_flag_reader,
        db_factory=lambda: sqlite3.connect(str(db_path)),
        bronze_root_resolver=lambda: tmp_path / "bronze",
    )

    with caplog.at_level(logging.INFO, logger="kairix.worker"):
        result = run_connector_sync_pipeline(deps)

    # Obsidian (flagless) still ran and indexed both notes.
    assert result.synced == 2, f"flagless obsidian sibling must still sync; got {result}"
    assert result.failed == 0

    gated_logs = [
        rec.getMessage()
        for rec in caplog.records
        if "sharepoint" in rec.getMessage().lower() and "gated" in rec.getMessage().lower()
    ]
    assert gated_logs, (
        f"expected a 'connector sharepoint gated off' INFO log; got {[r.getMessage() for r in caplog.records]}"
    )


@pytest.mark.unit
def test_multi_cc_pair_connector_yields_one_entry_per_pair(
    tmp_path: Path,
    caplog: pytest.LogCaptureFixture,
) -> None:
    """A single topology connector referenced by TWO cc_pairs yields two
    ingest entries (D2: one entry per cc_pair) — each carries the cc_pair
    name as its routing key.

    Proven cursor-independently: the connector kind is unresolvable
    (``does-not-exist``), so every entry fails at ``resolve_connector``
    and the per-entry try/except logs a failure naming the cc_pair. Two
    cc_pairs → two distinct failure logs (one per cc_pair name). If the
    enumeration yielded one entry per CONNECTOR instead of per cc_pair,
    only one of the two cc_pair names would ever appear.

    Sabotage proof: yield one entry per CONNECTOR (iterate
    ``parsed.connectors`` and emit a single entry per connector) → only
    one cc_pair name is logged, so the ``{"pair-a-cc", "pair-b-cc"}``
    assertion fails. Restored, both cc_pair names appear.
    """
    mapping: dict[str, Any] = {
        "topology": {
            "connectors": [
                {
                    "id": "shared-conn",
                    "kind": "does-not-exist",
                    "name": "Shared Connector",
                    "connector_specific_config": {},
                }
            ],
            "cc_pairs": [
                {"id": "pair-a", "connector": "shared-conn", "credential": None, "name": "pair-a-cc"},
                {"id": "pair-b", "connector": "shared-conn", "credential": None, "name": "pair-b-cc"},
            ],
        }
    }

    db_path = tmp_path / "index.sqlite"

    deps = ConnectorSyncDeps(
        disabled_fn=lambda: False,
        config_mapping_fn=lambda: mapping,
        db_factory=lambda: sqlite3.connect(str(db_path)),
        bronze_root_resolver=lambda: tmp_path / "bronze",
    )

    with caplog.at_level(logging.WARNING, logger="kairix.worker"):
        run_connector_sync_pipeline(deps)

    failed_pairs = {
        name
        for name in ("pair-a-cc", "pair-b-cc")
        if any(name in rec.getMessage() and "failed" in rec.getMessage() for rec in caplog.records)
    }
    assert failed_pairs == {"pair-a-cc", "pair-b-cc"}, (
        "each cc_pair must produce its own ingest entry (one entry per cc_pair, D2); "
        f"only these cc_pair names reached the per-entry loop: {failed_pairs}."
    )


# Marker substring identifying the operator-facing counter bump
# (``UPDATE topology_cc_pairs SET total_docs_indexed = ... WHERE id = ?``).
# Used by _CounterFailingConnection to fail ONLY that statement while every
# other write (bronze, silver, chunk-writer, cursor) commits normally.
_COUNTER_UPDATE_MARKER = "total_docs_indexed = total_docs_indexed + ?"


class _CounterFailingConnection(sqlite3.Connection):
    """A real ``sqlite3.Connection`` that raises on the counter-bump UPDATE only.

    Subclassing the real connection keeps the entire production pipeline
    (bronze store, silver processor, chunk writer, cursor store, dead-letter
    store) running against a live SQLite DB — F47-clean, no fakes at the
    storage seam. The ONLY divergence is that the ``total_docs_indexed``
    increment raises ``sqlite3.OperationalError``, so the public
    ``run_connector_sync_pipeline`` exercises the helper's best-effort
    ``except`` branch without any test reaching into a private name (F5).
    """

    def execute(self, sql: str, *args: Any) -> sqlite3.Cursor:
        if _COUNTER_UPDATE_MARKER in sql:
            raise sqlite3.OperationalError("simulated: topology_cc_pairs counter write failed")
        return super().execute(sql, *args)


@pytest.mark.integration
def test_counter_bump_db_failure_is_swallowed_and_loop_completes(
    tmp_path: Path,
    caplog: pytest.LogCaptureFixture,
) -> None:
    """C1 robustness through the public surface: a failure in the lifetime
    ``total_docs_indexed`` counter UPDATE must NOT abandon the sync.

    The counter bump runs in its own transaction AFTER the chunk writes have
    already committed; a failure there is best-effort — it logs a warning and
    the pipeline reports its real synced counters. A real batch indexes both
    notes (so the bump path is reached with ``indexed > 0`` and a resolved
    cc_pair id), but the connection raises on the counter UPDATE only. The
    aggregate result must still report ``synced == 2`` and ``failed == 0``,
    and a warning naming ``total_docs_indexed`` must be logged.

    Drives the helper's ``except Exception`` branch entirely through the
    public ``run_connector_sync_pipeline`` (no private-name import — F5).

    Sabotage proof: remove the ``try/except`` in
    ``_bump_cc_pair_total_docs_indexed`` (issue ``db.execute(...)`` bare); the
    ``OperationalError`` propagates out of ``run_connector_sync_pipeline`` and
    this test fails with an unhandled error. Restored, the bump failure is
    swallowed, the warning is logged, and ``synced == 2`` holds.
    """
    cc_pair_name = "obsidian-personal"
    vault = tmp_path / "vault"
    vault.mkdir()
    (vault / "alpha.md").write_text("# Alpha\n\nFirst note body content.\n", encoding="utf-8")
    (vault / "beta.md").write_text("# Beta\n\nSecond note body content.\n", encoding="utf-8")

    db_path = tmp_path / "index.sqlite"
    _seed_cc_pair_row(db_path, cc_pair_name=cc_pair_name)

    deps = ConnectorSyncDeps(
        disabled_fn=lambda: False,
        config_mapping_fn=lambda: _obsidian_topology(vault, cc_pair_name=cc_pair_name),
        db_factory=lambda: sqlite3.connect(str(db_path), factory=_CounterFailingConnection),
        bronze_root_resolver=lambda: tmp_path / "bronze",
    )

    with caplog.at_level(logging.WARNING, logger="kairix.worker"):
        result = run_connector_sync_pipeline(deps)

    # The batch succeeded; only the (best-effort) counter bump failed. The
    # pipeline must report its true synced counters and never propagate.
    assert result.synced == 2, (
        f"the batch indexed both notes; a counter-bump failure must not change synced. got {result}"
    )
    assert result.failed == 0, "a best-effort counter failure must not count as a connector failure"

    # The counter UPDATE failed, so the persisted value stays at its seeded 0.
    assert _total_docs_indexed(db_path, cc_pair_name=cc_pair_name) == 0

    warnings = [
        rec.getMessage()
        for rec in caplog.records
        if rec.levelno == logging.WARNING and "total_docs_indexed" in rec.getMessage()
    ]
    assert warnings, (
        "a counter-bump failure must emit a WARNING naming total_docs_indexed; "
        f"got {[r.getMessage() for r in caplog.records]}"
    )
